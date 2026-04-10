"""
Módulo de chat e utilitários relacionados.

Fornece:
- funções para extrair texto/imagens de arquivos (`extrair_texto`, `pdf_to_images`),
- integração com DuckDuckGo (`buscar_web`),
- `HistoricManager` para persistir conversas locais, e
- `Chat`, wrapper para interagir com modelos Ollama com suporte a RAG,
    imagens e ferramentas (functions/tools).

Observações:
- O arquivo contém um import suspeito `from turtle import pd` — isso parece ser
    um erro/typo (possivelmente pretendia importar `pandas as pd`). Não alterei
    a lógica; apenas documentei o ponto.
"""

import math
import os
import re
import sys
import tempfile
import zipfile
from turtle import pd
from httpx import delete
import ollama
import json
import base64

from collections.abc import Iterator
from email import message
from pathlib import Path

from typing import Any, Callable, Literal, Mapping, Sequence, List, Dict, Union, Optional
from ..generate.generate import Generate
from copy import deepcopy
from pydantic.json_schema import JsonSchemaValue
from datetime import date, datetime
from utils.utils import *

from ollama._types import (
  ChatResponse,
  Tool,
  ResponseError,
)

import logging
logging.getLogger("PyPDF2").setLevel(logging.ERROR)

from ddgs import DDGS


from pathlib import Path
from ollama import chat
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd

from ..rag.rag import RagIndex


_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}

def _extrair_imgs_zip(path: Path, media_dir: str) -> List[str]:
    imgs = []
    try:
        with zipfile.ZipFile(str(path), 'r') as z:
            for name in z.namelist():
                if name.startswith(media_dir) and Path(name).suffix.lower() in _IMAGE_EXTS:
                    try:
                        data = z.read(name)
                        imgs.append(base64.b64encode(data).decode("utf-8"))
                    except Exception:
                        pass
    except Exception:
        pass
    return imgs


def pdf_to_images(path: str | Path, dpi: int = 300) -> list[str]:
    """Converte cada página de um PDF em imagem PNG temporária.
    Requer PyMuPDF: pip install pymupdf
    Retorna lista de caminhos dos arquivos temporários criados.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF não está instalado. Execute: pip install pymupdf"
        ) from None

    p = Path(path)
    doc = fitz.open(str(p))
    tmp_files: list[str] = []
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    try:
        for page in doc:
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)  # type: ignore[union-attr]
            tmp = tempfile.NamedTemporaryFile(
                delete=False, suffix=".png", prefix="ollama_pdf_page_"
            )
            tmp_name = tmp.name
            tmp.close()  # Fecha antes de salvar (necessário no Windows — lock exclusivo)
            pix.save(tmp_name)
            tmp_files.append(tmp_name)
    finally:
        doc.close()
    return tmp_files


def extrair_texto(path: str) -> tuple[str, str, List[str]]:
    p = Path(path)
    ext = p.suffix.lower()

    if ext in {".png", ".jpg", ".jpeg", ".webp"}:
        return "image", str(p), []

    if ext in {".txt", ".md", ".py", ".json", ".xml", ".csv"}:
        return ext, p.read_text(encoding="utf-8", errors="ignore"), []

    if ext == ".pdf":
        reader = PdfReader(str(p))
        partes = []
        for page in reader.pages:
            try:
                partes.append(page.extract_text() or "")
            except Exception:
                partes.append("")
        texto = "\n".join(partes)
        imgs = []
        for page in reader.pages:
            try:
                page_images = page.images
            except Exception:
                continue
            for img in page_images:
                try:
                    imgs.append(base64.b64encode(img.data).decode("utf-8"))
                except Exception:
                    pass
        return ".pdf", texto, imgs

    if ext == ".docx":
        try:
            doc = Document(str(p))
            texto = "\n".join(par.text for par in doc.paragraphs if par.text.strip())
            imgs = _extrair_imgs_zip(p, "word/media/")
        except Exception as e:
            sanitized_path = sanitize_docx(p)
            doc = Document(str(sanitized_path))
            texto = "\n".join(par.text for par in doc.paragraphs if par.text.strip())
            imgs = _extrair_imgs_zip(Path(sanitized_path), "word/media/")
            os.remove(sanitized_path)
        return ".docx", texto, imgs

    if ext == ".xlsx":
        df = pd.read_excel(p, engine="openpyxl")
        imgs = _extrair_imgs_zip(p, "xl/media/")
        return ".xlsx", df.to_markdown(index=False), imgs

    raise ValueError(f"Tipo de arquivo não suportado: {ext}")

def _resize_image_file(path: str, max_px: int = 1600) -> None:
    """Redimensiona a imagem no arquivo `path` para que nenhuma dimensão exceda `max_px`.
    Salva em JPEG para reduzir ainda mais o tamanho. O arquivo original é substituído."""
    from PIL import Image
    img = Image.open(path)
    w, h = img.size
    if w <= max_px and h <= max_px:
        return
    scale = max_px / max(w, h)
    new_size = (int(w * scale), int(h * scale))
    img = img.resize(new_size, Image.LANCZOS) #type: ignore
    img.save(path, format="PNG", optimize=True)


def buscar_web(query: str, max_results: int=10 ) -> str:
    """Realiza uma busca na web usando DuckDuckGo e retorna os resultados como uma string.
    Args:
        query (str): A consulta de busca.
        max_results (int): O número máximo de resultados a serem retornados.
    """
    for _ in range(3):
        with DDGS() as ddgs:
            resultados = list(ddgs.text(query, max_results=max_results))
        if resultados:
            return str(resultados)
    return "No search results found."


class HistoricManager():
    """Gerencia histórico de conversas salvo em disco.

    O histórico é mantido em um arquivo JSON com título → lista de mensagens.
    Exemplo de uso::

        hm = HistoricManager()
        hm.register_chat('minha_conversa', [{'role':'user','content':'Olá'}])
        print(hm.list_chats())
        chat = hm.get_chat('minha_conversa')

    O construtor aceita um caminho para o arquivo JSON ou para um diretório
    (neste caso, cria `chat.json` dentro do diretório).
    """
    @property
    def data(self) -> Dict[str, Any]:
        return self.__data
    
    def __init__(self, path:Path|str=Path.cwd().joinpath("historic", "chat.json")):
        if isinstance(path, str):
            path = Path(path)
        
        if path.is_file():
            path.parent.mkdir(parents=True, exist_ok=True)
            if not path.name.endswith(".json"):
                path = path.with_suffix(".json")
            
        elif path.is_dir():
            path.mkdir(parents=True, exist_ok=True)
            path = path.joinpath("chat.json")
            
        if not path.exists():
            with open(path, "w", encoding="utf-8") as f:
                json.dump({}, f, ensure_ascii=False, indent=4)
                
        self.__data: Dict[str, Any] = {}
        self.__path = path
                
    def __save(self, data: dict):
        with open(self.__path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        
    def __load(self) -> dict:
        with open(self.__path, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.__data = data
        return data
    
    def register_chat(self, title: str, chat: List[dict]):
        data = self.__load()
        data[title] = chat
        self.__save(data)
        
    def list_chats(self) -> List[str]:
        data = self.__load()
        return list(data.keys())
    
    def get_chat(self, title: str) -> List[dict]:
        data = self.__load()
        return data.get(title, [])
    
    def delete_chat(self, title: str):
        data = self.__load()
        if title in data:
            del data[title]
            self.__save(data)
    
    def clear_all_chats(self):
        self.__save({})

class Chat():
    """Wrapper de alto nível para conversar com modelos via Ollama.

    A classe encapsula lógica para:
    - gerenciar histórico (`historicchat`),
    - detectar contexto/token limits do modelo e ajustar `num_ctx`,
    - incluir arquivos (texto, pdf, docx, xlsx) e imagens, e
    - injetar contexto RAG obtido de índices gerados por :class:`RagIndex`.

    Exemplo mínimo::

        chat = Chat(model='meu-modelo')
        resp = chat.chat('Olá')

    Observações:
    - Muitos parâmetros avançados existem no método `chat()` (temperature, think,
      rag_paths, etc.). Consulte a docstring desse método para detalhes.
    """
    # @property
    # def model(self) -> str:
    #     return self.__model
        
    def __init__(
        self, 
        model:str, 
        chat_sistem:str = "",
        host: Literal["http://localhost:11434"]| str | None = None, 
        **kwargs
    ) -> None:    
          
        self.model = model
        
        self.historicchat: List[dict] = []
        
        self.__client = ollama.Client(host, **kwargs)
        self.__context_size_cache: dict[str, int | None] = {}

        self.chat_sistem = chat_sistem
        # if self.__system:
        #     self.__historychat.append({"role": "system", "content": system})
            
            
        
    def restart_chat(self, system: str = "") -> None:
        self.historicchat = []
        if system:
            self.historicchat.append({"role": "system", "content": system})

    def _get_context_size(self) -> int | None:
        if self.model in self.__context_size_cache:
            return self.__context_size_cache[self.model]
        try:
            info = self.__client.show(self.model)
            ctx = None

            # Prioridade 1: num_ctx em parameters (ex: "num_ctx 32768")
            parameters = getattr(info, "parameters", None) or ""
            for line in parameters.splitlines():
                parts = line.strip().split()
                if len(parts) == 2 and parts[0].lower() == "num_ctx":
                    try:
                        ctx = int(parts[1])
                    except ValueError:
                        pass
                    break

            # Prioridade 2: PARAMETER num_ctx no modelfile completo
            if ctx is None:
                modelfile = getattr(info, "modelfile", None) or ""
                for line in modelfile.splitlines():
                    parts = line.strip().split()
                    if len(parts) == 3 and parts[0].upper() == "PARAMETER" and parts[1].lower() == "num_ctx":
                        try:
                            ctx = int(parts[2])
                        except ValueError:
                            pass
                        break

            # Prioridade 3: context_length nos metadados GGUF
            if ctx is None:
                modelinfo = getattr(info, "modelinfo", None) or {}
                raw = next(
                    (v for k, v in modelinfo.items() if k.endswith(".context_length")),
                    None,
                )
                if raw is not None:
                    try:
                        ctx = int(raw)
                    except (ValueError, TypeError):
                        pass

            logging.getLogger(__name__).debug(
                f"[context_size] model='{self.model}' ctx_size={ctx}"
            )
            self.__context_size_cache[self.model] = ctx
            return ctx
        except Exception as e:
            logging.getLogger(__name__).debug(
                f"[context_size] model='{self.model}' show() falhou: {e} — verificação ignorada"
            )
            self.__context_size_cache[self.model] = None
            return None

    @staticmethod
    def _estimate_tokens(messages: list) -> int:
        total = 0
        for msg in messages:
            content = msg.get("content", "")
            if isinstance(content, str):
                total += len(content) // 4
            images = msg.get("images", [])
            total += len(images) * 85
        return total

    @staticmethod
    def convert_images_to_base64(images_paths: List[str]) -> List[dict]:
        images = []
        for image_path in images_paths:
            path = Path(image_path)
            if not path.is_file() or not path.exists():
                raise ValueError(f"Image path '{image_path}' is not a valid file.")
            # if not path.suffix.lower() in [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".pdf"]:
            #     raise ValueError(f"Image path '{image_path}' is not a supported image format.")
            
            images.append(base64.b64encode(path.read_bytes()).decode("utf-8"))
        return images
        
    def build_rag(
        self,
        files_paths: List[str],
        save_path: str,
        embedding_model: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        format: Literal['json', 'npz'] = 'json',
    ) -> RagIndex:
        """Constrói e salva um índice RAG a partir de arquivos.

        Args:
            files_paths: Arquivos de origem (txt, md, pdf, docx, xlsx, etc.).
            save_path: Caminho de destino do arquivo gerado (ex: ``'meu_rag.json'``).
            embedding_model: Modelo de embedding Ollama (ex: ``'nomic-embed-text'``).
            chunk_size: Tamanho máximo de cada chunk em caracteres.
            chunk_overlap: Sobreposição em caracteres entre chunks consecutivos.
            format: Formato do arquivo — ``'json'`` ou ``'npz'``.

        Returns:
            :class:`~agente_ollama.rag.rag.RagIndex` carregado em memória.
        """
        host = getattr(self.__client, '_host', None) or getattr(self.__client, 'host', None)
        return RagIndex.build(
            files_paths=files_paths,
            save_path=save_path,
            embedding_model=embedding_model,
            host=host,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            format=format,
        )

    def chat(self, 
            prompt: str,
            web_search: Optional[bool]=False,
            system:str="",
            format:Literal['', 'json']='',
            files_paths:List[str]=[],
            temperature:float|None=None,
            top_p:float|None=None,
            seed:int|None=None,
            num_predictions:int|None=None,
            think: Optional[Union[bool, Literal['low', 'medium', 'high']]] = None,
            date:datetime|None = None,
            remove_historic:bool=True,
            remove_date:bool=True,
            
            logprobs: Optional[bool] = None,
            top_logprobs: Optional[int] = None,
            keep_alive: Optional[Union[float, str]] = None,
            truncate_on_overflow: bool = True,
            pdf_images: bool = False,
            pdf_as_images: bool = False,
            pdf_dpi: int = 300,
            debug_print_messages: bool = False,
            rag_paths: List[str] = [],
            rag_embedding_model: str = "",
            rag_top_k: int = 5,
            #tools: Optional[Sequence[Union[Mapping[str, Any], Tool, Callable]]] = None,
            
        ) -> ChatResponse:
        if isinstance(date, datetime):
            _date = date
        else:
            _date = datetime.now()
        
        options = {}
        if temperature is not None:
            options['temperature'] = temperature
        if top_p is not None:
            options['top_p'] = top_p
        if seed is not None:
            options['seed'] = seed
        if num_predictions is not None:
            options['num_predictions'] = num_predictions
        
        
        #_images = self.convert_images_to_base64(images_paths) if images_paths else []
        
        #create historic   
        if remove_historic:       
            message = []
        else:
            message = deepcopy(self.historicchat)
            
        #system unifique
        system_temp = [self.chat_sistem] if self.chat_sistem else []
        for msg in message:
            if msg['role'] == 'system':
                sys_temp = msg['content'].strip().split("\n")
                for s in sys_temp:
                    system_temp.append(s)
                message.remove(msg)
        
        if not system in system_temp:
            system_temp.append(system)
            
        system_temp = list(set(system_temp))
        system_temp = [s for s in system_temp if s.strip()]
                    
        message.insert(0,{"role": "system", "content": ("\n".join(system_temp)).strip()})
        
        #import pdb;pdb.set_trace()
         
        if not remove_date:
            date_str = _date.strftime("%Y-%m-%d %H:%M:%S")
            sys_content = message[0]['content'] if message and message[0]['role'] == 'system' else ""
            date_line = f"Data/hora atual: {date_str}"
            if message and message[0]['role'] == 'system':
                message[0]['content'] = (sys_content + "\n" + date_line).strip() if sys_content else date_line
            else:
                message.insert(0, {"role": "system", "content": date_line})
        
            
        
        
        _images = []
        _temp_img_files: list[str] = []
        texts = ""
        for file_path in files_paths:
            if pdf_as_images and Path(file_path).suffix.lower() == ".pdf":
                try:
                    page_imgs = pdf_to_images(file_path, dpi=pdf_dpi)
                    for img_path in page_imgs:
                        try:
                            _resize_image_file(img_path)
                        except Exception as resize_err:
                            logging.getLogger(__name__).debug(
                                f"[resize] falha ao redimensionar '{img_path}': {resize_err}"
                            )
                    _temp_img_files.extend(page_imgs)
                    _images.extend(page_imgs)
                except Exception as e:
                    logging.getLogger(__name__).warning(
                        f"[pdf_as_images] falha ao converter '{file_path}': {e}"
                    )
                continue

            _type, content, embedded_imgs = extrair_texto(file_path)
            if _type == "image":
                _images.append(content)
            else:
                texts += f"'{_type}':\n{content}\n\n"
            if pdf_images and _type == ".pdf":
                for img_b64 in embedded_imgs:
                    try:
                        img_data = base64.b64decode(img_b64)
                        tmp = tempfile.NamedTemporaryFile(
                            delete=False, suffix=".png", prefix="ollama_pdf_img_"
                        )
                        tmp.write(img_data)
                        tmp.close()
                        _images.append(tmp.name)
                        _temp_img_files.append(tmp.name)
                    except Exception:
                        pass
            elif _type != ".pdf":
                _images.extend(embedded_imgs)

        # RAG: buscar chunks relevantes e injetar no system
        if rag_paths:
            rag_chunks: List[str] = []
            _rag_emb_model = rag_embedding_model
            for rag_path in rag_paths:
                try:
                    _idx = RagIndex.load(rag_path)
                    _emb_model = _rag_emb_model or _idx._model
                    _host = getattr(self.__client, '_host', None) or getattr(self.__client, 'host', None)
                    results = _idx.search(
                        query=prompt,
                        embedding_model=_emb_model,
                        top_k=rag_top_k,
                        host=_host,
                    )
                    rag_chunks.extend(results)
                except Exception as _rag_err:
                    logging.getLogger(__name__).warning(
                        f"[RAG] falha ao usar índice '{rag_path}': {_rag_err}"
                    )
            if rag_chunks:
                rag_context = (
                    "Use as informações abaixo para responder ao usuário. "
                    "Não mencione que recebeu contexto externo, arquivos ou base de conhecimento — "
                    "utilize as informações de forma natural e transparente.\n\n"
                    + "\n---\n".join(rag_chunks)
                )
                sys_msg = message[0] if message and message[0]['role'] == 'system' else None
                if sys_msg:
                    sys_msg['content'] = (sys_msg['content'] + "\n" + rag_context).strip()
                else:
                    message.insert(0, {"role": "system", "content": rag_context})

        original_prompt = prompt
        if texts:
            prompt = f"{prompt}\n\nfiles:\n{texts}"

        role:dict = {"role": "user", "content": prompt}

        if _images:
           role['images'] = _images[:8]

        message.append(role)

        # Verificação de tokens
        _RESPONSE_RESERVE = 512
        ctx_size = self._get_context_size()
        if ctx_size is not None:
            limit = ctx_size - _RESPONSE_RESERVE
            estimated = self._estimate_tokens(message)
            if estimated > limit:
                if not truncate_on_overflow:
                    raise ValueError(
                        f"Tokens estimados ({estimated}) excedem o limite do modelo "
                        f"'{self.model}' ({ctx_size} tokens, reserva de {_RESPONSE_RESERVE})."
                    )
                # Truncar texts progressivamente
                if texts:
                    _block = 1000
                    while texts and self._estimate_tokens(message) > limit:
                        texts = texts[:-_block]
                        new_prompt = f"{original_prompt}\n\nfiles:\n{texts}" if texts else original_prompt
                        role["content"] = new_prompt

            # Define num_ctx dinamicamente com o mínimo necessário para evitar
            # alocação desnecessária de memória no servidor Ollama
            estimated_final = self._estimate_tokens(message)
            optimal_ctx = min(estimated_final + _RESPONSE_RESERVE, ctx_size)
            # Arredonda para próxima potência de 2 para melhor alinhamento de memória
            optimal_ctx = 2 ** math.ceil(math.log2(max(optimal_ctx, 512)))
            optimal_ctx = min(optimal_ctx, ctx_size)
            options['num_ctx'] = optimal_ctx
            logging.getLogger(__name__).debug(
                f"[num_ctx] estimated={estimated_final} → num_ctx={optimal_ctx} (max={ctx_size})"
            )

        if web_search:
            try:
                #web_search
                search = self.web_search(
                    query=prompt
                )
                if search:
                    for item in search:
                        message.append(item)
            except Exception as err:
                pass
            
        #print(f"\n\n{message=}\n\n")

        if debug_print_messages:
            import pprint
            print("\n===== DEBUG: messages enviadas ao LLM =====")
            pprint.pprint(message)
            print("=========================================\n")

        def _do_chat(msgs: list, opts: dict) -> ChatResponse:
            return self.__client.chat(
                model=self.model,
                messages=msgs,
                stream=False,
                format=format,
                think=think,
                logprobs=logprobs,
                top_logprobs=top_logprobs,
                keep_alive=keep_alive,
                options=opts if opts else None,
            )

        ##########     CHAT
        try:
            try:
                res = _do_chat(message, options)
            except ResponseError as e:
                if e.status_code == 400 and "length limit exceeded" in e.error:
                    # Request body muito grande — reduz imagens pela metade e tenta novamente
                    logging.getLogger(__name__).debug(
                        f"[chat] 400 length limit exceeded, reduzindo imagens para '{self.model}'"
                    )
                    for msg in message:
                        imgs = msg.get("images")
                        if imgs:
                            msg["images"] = imgs[: max(1, len(imgs) // 2)]
                    res = _do_chat(message, options)
                elif (
                    e.status_code == 500
                    and "model failed to load" not in e.error
                    and "system memory" not in e.error
                    and any(msg.get("images") for msg in message)
                ):
                    # Retry sem imagens — imagens embutidas podem travar o model runner
                    logging.getLogger(__name__).debug(
                        f"[chat] 500 com imagens, retentando sem imagens para '{self.model}'"
                    )
                    for msg in message:
                        msg.pop("images", None)
                    # Recalcula num_ctx sem o peso das imagens
                    if ctx_size is not None:
                        estimated_retry = self._estimate_tokens(message)
                        optimal_retry = min(estimated_retry + _RESPONSE_RESERVE, ctx_size)
                        optimal_retry = 2 ** math.ceil(math.log2(max(optimal_retry, 512)))
                        optimal_retry = min(optimal_retry, ctx_size)
                        options['num_ctx'] = optimal_retry
                    res = _do_chat(message, options)
                else:
                    raise
        finally:
            for _tmp in _temp_img_files:
                try:
                    os.remove(_tmp)
                except Exception:
                    pass
        ##########
        
        if res.message.content:
            message.append({"role": "assistant", "content": res.message.content})
        else:
            raise Exception(f"{res}")
        
        _message = []
        for msg in message:
            if msg['role'] == 'tool':
                #message.remove(msg)
                continue
            if 'tool_calls' in msg:
                #message.remove(msg)
                continue
            #import pdb;pdb.set_trace()
            _message.append(msg)
                
        #print(message)
        if not remove_historic:
            self.historicchat = _message
        
        return res
    
    def web_search(self, query:str) -> List:
        tools_message = []
        tools_message.append({"role": "system", "content": "You are a helpful assistant with access to a web search tool. When the user provides a query, you will use the 'buscar_web' tool to search the web and return relevant information."})
        
        tools_message.append({"role": "user", "content": query})
        
        tools = [buscar_web]
        tools_map = {t.__name__: t for t in tools if callable(t)}
        res = self.__client.chat(
            model=self.model, 
            messages=tools_message,
            stream=False,
            tools=tools,
        )
        ##########
        
        #import pdb;pdb.set_trace()
        tools_return = []
        if res.message.tool_calls:
            for tool in res.message.tool_calls:
                function = tool.function.name
                args = tool.function.arguments
                                        
                # procura a função nas tools passadas, senão em globals()
                func = tools_map.get(function) or globals().get(function)
                if func is None:
                    raise NameError(f"Tool function '{function}' não encontrada")

                # chamar adequadamente conforme tipo de args
                if isinstance(args, dict):
                    result_function = func(**args)
                elif isinstance(args, (list, tuple)):
                    result_function = func(*args)
                else:
                    result_function = func(args) # type: ignore
                
                # tools_message.append({"role": "assistant", "content": result_function})
                tools_return.append(
                    {
                        "role": "assistant",
                        "tool_calls": [
                            {
                                'type': 'function',
                                'function': {
                                    'name': function,
                                    'arguments': args
                                },
                            }
                        ]
                    }
                )
                
                tools_return.append({"role": "tool", "tool_name": function, "content": result_function})
                        
        return tools_return
        
    def functions_tools(
        self, 
        prompt:str,
        tools:List[Any],
        think: Optional[Union[bool, Literal['low', 'medium', 'high']]] = None,
        system:str = "",
    ) -> ChatResponse:
        tools_message = []
        if system:
            tools_message.append({"role": "system", "content": system})
        
        tools_message.append({"role": "user", "content": prompt})
        
        #for _ in range(len())
        ##########     CHAT
        tools_map = {t.__name__: t for t in tools if callable(t)}
        res = self.__client.chat(
            model=self.model, 
            messages=tools_message,
            stream=False,
            think=think,
            tools=tools,
            
        )
        ##########
        
        #import pdb;pdb.set_trace()
        
        if res.message.tool_calls:
            for tool in res.message.tool_calls:
                function = tool.function.name
                args = tool.function.arguments
                                
                
                # procura a função nas tools passadas, senão em globals()
                func = tools_map.get(function) or globals().get(function)
                if func is None:
                    raise NameError(f"Tool function '{function}' não encontrada")

                # chamar adequadamente conforme tipo de args
                if isinstance(args, dict):
                    result_function = func(**args)
                elif isinstance(args, (list, tuple)):
                    result_function = func(*args)
                else:
                    result_function = func(args)                
                
                # tools_message.append({"role": "assistant", "content": result_function})
                tools_message.append(
                    {
                        "role": "assistant",
                        "tool_calls": [
                            {
                                'type': 'function',
                                'function': {
                                    'name': function,
                                    'arguments': args
                                },
                            }
                        ]
                    }
                )
                
                
                tools_message.append({"role": "tool", "tool_name": function, "content": result_function})
                
                
        res = self.__client.chat(
            model=self.model, 
            messages=tools_message,
            stream=False,
            think=think,
        )
        
        tools_message.append({"role": "assistant", "content": res.message.content})
        
        #print(f"\n\n{tools_message}\n\n")
        return res
        
        

if __name__ == "__main__":
    pass
